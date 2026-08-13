#!/usr/bin/env python3
"""
Step 4 — Unstructured document classification for the E-Rakshak ingestion pipeline.

Targets non-tabular files left in uploads/raw/ after Steps 2–3
(PDF, Word, images). Extracts text from the first page (or first image via
optional OCR) and classifies with weighted keyword rules:

  fir   — First Information Report / police complaint language
  kyc   — Aadhaar / PAN / ID / subscriber KYC
  cctv  — camera / DVR / footage logs
  bank  — bank statement PDFs (not caught by Step 3 tabular rules)
  other — readable but no strong match (case diary, notices, etc.)

Unreadable / empty extractions → uploads/quarantine/

Usage:
    python ingestion_p/classify_unstructured.py
    python ingestion_p/classify_unstructured.py --dry-run
"""

from __future__ import annotations

import argparse
import json
import logging
import re
import shutil
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path

try:
    from setup_staging import UPLOADS_ROOT, create_staging_layout
except ImportError:  # pragma: no cover
    from ingestion_p.setup_staging import UPLOADS_ROOT, create_staging_layout

logger = logging.getLogger(__name__)

# Extensions handled in this step (tabular already routed in Step 3)
DOC_EXTENSIONS = {".pdf", ".doc", ".docx"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp", ".gif"}
UNSTRUCTURED_EXTENSIONS = DOC_EXTENSIONS | IMAGE_EXTENSIONS | {".html", ".htm", ".eml", ".txt"}

# Max characters kept from first-page extract (keeps classification fast)
MAX_TEXT_CHARS = 8000

MIN_SCORE = 4
MIN_MARGIN = 2

# (keyword_or_phrase, weight) — matched case-insensitively on extracted text
# Override / extend via ingestion_p/keyword_rules.json:
#   {"fir": [["fir no", 6], ...], "kyc": [...], ...}
CATEGORY_KEYWORDS: dict[str, list[tuple[str, int]]] = {
    "fir": [
        ("first information report", 8),
        ("fir no", 6),
        ("fir number", 6),
        ("fir/", 4),
        (" police station", 3),
        ("complainant", 4),
        ("accused", 3),
        ("under section", 4),
        ("u/s", 3),
        ("ipc", 3),
        ("it act", 3),
        ("cyber crime", 3),
        ("nccrp", 4),
        ("investigation", 1),
        ("cognizable", 3),
        ("registered the fir", 5),
    ],
    "kyc": [
        ("aadhaar", 6),
        ("aadhar", 6),
        ("permanent account number", 5),
        (" pan ", 3),
        ("pan card", 5),
        ("ckyc", 5),
        ("know your customer", 6),
        ("identity proof", 4),
        ("address proof", 4),
        ("voter id", 4),
        ("passport no", 4),
        ("driving licence", 4),
        ("driving license", 4),
        ("subscriber information", 5),
        ("customer id", 2),
        ("photograph", 2),
        ("date of birth", 2),
    ],
    "cctv": [
        ("cctv", 7),
        ("camera", 3),
        ("dvr", 4),
        ("nvr", 4),
        ("footage", 5),
        ("recording", 3),
        ("surveillance", 4),
        ("channel no", 3),
        ("video clip", 4),
        ("hard disk", 2),
    ],
    "bank": [
        ("account statement", 6),
        ("bank statement", 7),
        ("statement of account", 6),
        ("opening balance", 4),
        ("closing balance", 4),
        ("withdrawal", 3),
        ("deposit", 2),
        ("ifsc", 4),
        ("account number", 3),
        ("a/c no", 3),
        ("transaction", 2),
        ("upi", 3),
        ("imps", 3),
        ("neft", 3),
        ("rtgs", 3),
        ("debit", 2),
        ("credit", 2),
        ("branch", 1),
    ],
}


def _load_keyword_overrides() -> None:
    """Merge optional keyword_rules.json into CATEGORY_KEYWORDS (additive)."""
    candidates = [
        Path(__file__).resolve().parent / "keyword_rules.json",
        Path(__file__).resolve().parent / "config" / "keyword_rules.json",
    ]
    for path in candidates:
        if not path.is_file():
            continue
        try:
            raw = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError) as exc:
            logger.warning("Could not load keyword rules from %s: %s", path, exc)
            continue
        if not isinstance(raw, dict):
            continue
        for cat, rules in raw.items():
            if not isinstance(rules, list):
                continue
            bucket = CATEGORY_KEYWORDS.setdefault(str(cat).lower(), [])
            existing = {(k.lower(), int(w)) for k, w in bucket}
            for item in rules:
                if not isinstance(item, (list, tuple)) or len(item) < 2:
                    continue
                kw, weight = str(item[0]), int(item[1])
                if (kw.lower(), weight) not in existing:
                    bucket.append((kw, weight))
                    existing.add((kw.lower(), weight))
        logger.info("Loaded keyword overrides from %s", path)
        break


_load_keyword_overrides()


@dataclass
class DocClassificationResult:
    path: str
    original_name: str
    category: str
    confidence: float
    scores: dict[str, int]
    matched_keywords: list[str]
    text_preview: str
    status: str  # classified | quarantined | skipped | failed
    reason: str
    destination: str | None = None


@dataclass
class DocClassificationReport:
    scanned_at: str
    raw_dir: str
    processed_dir: str
    quarantine_dir: str
    total_seen: int = 0
    considered: int = 0
    classified: int = 0
    other: int = 0
    quarantined: int = 0
    skipped: int = 0
    failed: int = 0
    results: list[DocClassificationResult] = field(default_factory=list)


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip().lower()


def extract_pdf_first_page(path: Path) -> str:
    import pdfplumber

    with pdfplumber.open(path) as pdf:
        if not pdf.pages:
            return ""
        page = pdf.pages[0]
        return page.extract_text() or ""


def extract_docx_first_page(path: Path) -> str:
    """
    Word has no hard page break API here — take leading paragraphs
    until we roughly hit one page of text (~3000 chars).
    """
    import docx as python_docx

    doc = python_docx.Document(path)
    chunks: list[str] = []
    total = 0
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if not t:
            continue
        chunks.append(t)
        total += len(t)
        if total >= 3000:
            break
    # Also pull first table cells (FIRs / KYC often use tables)
    for table in doc.tables[:2]:
        for row in table.rows[:15]:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
                total += sum(len(c) for c in cells)
            if total >= 4000:
                break
        if total >= 4000:
            break
    return "\n".join(chunks)


def extract_doc_legacy(path: Path) -> str:
    """
    Best-effort for old .doc: pull printable ASCII/UTF-16LE runs.
    Avoids requiring MS Word / antiword.
    """
    data = path.read_bytes()
    # UTF-16LE runs (Word often stores text this way)
    try:
        utf16 = data.decode("utf-16-le", errors="ignore")
        utf16_clean = "".join(ch if ch.isprintable() or ch in "\n\r\t" else " " for ch in utf16)
        utf16_clean = re.sub(r" {2,}", " ", utf16_clean)
    except Exception:  # noqa: BLE001
        utf16_clean = ""
    ascii_chunks = re.findall(rb"[\x20-\x7e]{6,}", data)
    ascii_text = " ".join(c.decode("ascii", errors="ignore") for c in ascii_chunks[:200])
    combined = f"{utf16_clean}\n{ascii_text}"
    return combined[:MAX_TEXT_CHARS]


def extract_image_ocr(path: Path) -> str:
    """Optional OCR — requires pytesseract + system Tesseract binary."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.warning("OCR unavailable (install pytesseract + Pillow); skipping image text for %s", path.name)
        return ""

    try:
        with Image.open(path) as img:
            # First "page" of multi-frame TIFF
            text = pytesseract.image_to_string(img)
            return text or ""
    except Exception as exc:  # noqa: BLE001
        logger.warning("OCR failed for %s: %s", path.name, exc)
        return ""


def extract_plain_text(path: Path) -> str:
    for enc in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=enc)[:MAX_TEXT_CHARS]
        except UnicodeDecodeError:
            continue
    return ""


def extract_first_page_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf_first_page(path)
    if ext == ".docx":
        return extract_docx_first_page(path)
    if ext == ".doc":
        return extract_doc_legacy(path)
    if ext in IMAGE_EXTENSIONS:
        return extract_image_ocr(path)
    if ext in {".txt", ".html", ".htm", ".eml"}:
        return extract_plain_text(path)
    return ""


def score_text(text: str) -> tuple[dict[str, int], dict[str, list[str]]]:
    norm = normalize_text(text)
    # Pad so short-token patterns like " pan " can match edges
    padded = f" {norm} "
    scores = {cat: 0 for cat in CATEGORY_KEYWORDS}
    matched: dict[str, list[str]] = {cat: [] for cat in CATEGORY_KEYWORDS}

    for category, rules in CATEGORY_KEYWORDS.items():
        for phrase, weight in rules:
            needle = phrase.lower()
            if needle in padded or needle.strip() in padded:
                scores[category] += weight
                matched[category].append(phrase.strip())
    return scores, matched


def pick_category(scores: dict[str, int], matched: dict[str, list[str]]) -> tuple[str, float, list[str], str]:
    ranked = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
    best, best_score = ranked[0]
    second = ranked[1][1] if len(ranked) > 1 else 0

    if best_score < MIN_SCORE:
        return "other", 0.0, [], f"no strong keyword match (best={best}:{best_score})"

    if best_score - second < MIN_MARGIN and second >= MIN_SCORE:
        return "other", 0.2, matched.get(best, []), f"ambiguous scores={scores}"

    confidence = min(1.0, best_score / 16.0)
    return best, confidence, matched.get(best, []), f"matched {best} (score={best_score})"


def _unique_destination(dest_dir: Path, filename: str) -> Path:
    candidate = dest_dir / filename
    if not candidate.exists():
        return candidate
    stem = Path(filename).stem
    suffix = Path(filename).suffix
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S%f")
    return dest_dir / f"{stem}__{stamp}{suffix}"


def _move_with_sidecar(src: Path, dest: Path, payload: dict) -> None:
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.move(str(src), str(dest))
    meta = dest.with_suffix(dest.suffix + ".classify.json")
    meta.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def classify_unstructured_file(path: Path) -> DocClassificationResult:
    name = path.name
    try:
        text = extract_first_page_text(path)
        preview = re.sub(r"\s+", " ", text).strip()[:240]

        if not text or not text.strip():
            return DocClassificationResult(
                path=str(path),
                original_name=name,
                category="quarantine",
                confidence=0.0,
                scores={},
                matched_keywords=[],
                text_preview="",
                status="quarantined",
                reason="no extractable text on first page (scan/OCR may be required)",
            )

        scores, matched = score_text(text)
        category, confidence, hits, reason = pick_category(scores, matched)
        status = "classified"
        return DocClassificationResult(
            path=str(path),
            original_name=name,
            category=category,
            confidence=confidence,
            scores=scores,
            matched_keywords=hits,
            text_preview=preview,
            status=status,
            reason=reason,
        )
    except Exception as exc:  # noqa: BLE001
        return DocClassificationResult(
            path=str(path),
            original_name=name,
            category="quarantine",
            confidence=0.0,
            scores={},
            matched_keywords=[],
            text_preview="",
            status="failed",
            reason=f"{type(exc).__name__}: {exc}",
        )


def classify_raw_unstructured(
    raw_dir: Path | None = None,
    processed_dir: Path | None = None,
    quarantine_dir: Path | None = None,
    *,
    move_files: bool = True,
) -> DocClassificationReport:
    """
    Classify unstructured files remaining in uploads/raw/ after Steps 2–3.
    """
    create_staging_layout()

    raw = Path(raw_dir) if raw_dir else UPLOADS_ROOT / "raw"
    processed = Path(processed_dir) if processed_dir else UPLOADS_ROOT / "processed"
    quarantine = Path(quarantine_dir) if quarantine_dir else UPLOADS_ROOT / "quarantine"

    for d in (raw, processed, quarantine):
        d.mkdir(parents=True, exist_ok=True)
    for sub in ("fir", "kyc", "cctv", "bank", "other"):
        (processed / sub).mkdir(parents=True, exist_ok=True)

    report = DocClassificationReport(
        scanned_at=datetime.now(timezone.utc).isoformat(),
        raw_dir=str(raw),
        processed_dir=str(processed),
        quarantine_dir=str(quarantine),
    )

    files = sorted(p for p in raw.iterdir() if p.is_file() and p.name != ".gitkeep")
    report.total_seen = len(files)

    for path in files:
        ext = path.suffix.lower()
        if ext not in UNSTRUCTURED_EXTENSIONS:
            report.skipped += 1
            report.results.append(
                DocClassificationResult(
                    path=str(path),
                    original_name=path.name,
                    category="",
                    confidence=0.0,
                    scores={},
                    matched_keywords=[],
                    text_preview="",
                    status="skipped",
                    reason=f"unsupported for Step 4 (extension {ext})",
                )
            )
            continue

        report.considered += 1
        result = classify_unstructured_file(path)

        if result.status in {"quarantined", "failed"} or result.category == "quarantine":
            report.quarantined += 1
            if result.status == "failed":
                report.failed += 1
            dest_dir = quarantine
            result.category = "quarantine"
            result.status = "quarantined" if result.status != "failed" else result.status
        else:
            if result.category == "other":
                report.other += 1
            else:
                report.classified += 1
            dest_dir = processed / result.category

        dest = _unique_destination(dest_dir, path.name)
        if move_files:
            _move_with_sidecar(
                path,
                dest,
                {
                    "original_name": path.name,
                    "category": result.category,
                    "confidence": result.confidence,
                    "scores": result.scores,
                    "matched_keywords": result.matched_keywords,
                    "text_preview": result.text_preview,
                    "reason": result.reason,
                    "classified_at": datetime.now(timezone.utc).isoformat(),
                },
            )
            result.destination = str(dest)
        else:
            result.destination = str(dest)

        report.results.append(result)

    return report


def save_report(report: DocClassificationReport, path: Path | None = None) -> Path:
    out = path or (UPLOADS_ROOT / "unstructured_classification_report.json")
    payload = {
        **{k: v for k, v in asdict(report).items() if k != "results"},
        "results": [asdict(r) for r in report.results],
    }
    out.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return out


def _print_report(report: DocClassificationReport) -> None:
    print("=== Step 4: Unstructured Document Classification ===")
    print(f"raw         : {report.raw_dir}")
    print(f"processed   : {report.processed_dir}")
    print(f"quarantine  : {report.quarantine_dir}")
    print(f"seen        : {report.total_seen}")
    print(f"considered  : {report.considered}")
    print(f"classified  : {report.classified}")
    print(f"other       : {report.other}")
    print(f"quarantined : {report.quarantined}")
    print(f"skipped     : {report.skipped}")
    print()
    for r in report.results:
        if r.status == "skipped":
            print(f"  [SKIP] {r.original_name} — {r.reason}")
            continue
        label = (r.category or "?").upper()
        print(f"  [{label:12}] {r.original_name}  conf={r.confidence:.2f}")
        print(f"               keywords={r.matched_keywords[:8]}")
        print(f"               reason={r.reason}")
        if r.text_preview:
            print(f"               preview={r.text_preview[:120]}...")
        if r.destination:
            print(f"               -> {r.destination}")


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
    parser = argparse.ArgumentParser(
        description="Classify unstructured docs left in uploads/raw/"
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Classify and report only; do not move files",
    )
    args = parser.parse_args()
    report = classify_raw_unstructured(move_files=not args.dry_run)
    _print_report(report)
    out = save_report(report)
    print(f"\nReport saved: {out}")
    if args.dry_run:
        print("(dry-run: no files were moved)")


if __name__ == "__main__":
    main()
